"""
Build the conference paper (IEEE A4 template) from the project's real results.

Reuses the styles of the supplied IEEE template (paper/ieee_template.docx, a
Word-readable conversion of conference-template-a4.docx) and fills it with the
paper content. Figures are the actual plots in results/; every number quoted in
the tables is taken from the committed results/*.json files. Run:

    .venv/bin/python paper/build_paper.py
        → paper/RL_Adaptive_Routing.docx
"""

import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(__file__)
ROOT = os.path.dirname(HERE)
RESULTS = os.path.join(ROOT, "results")
TEMPLATE = os.path.join(HERE, "ieee_template.docx")
OUT = os.path.join(HERE, "RL_Adaptive_Routing.docx")


# ── helpers ─────────────────────────────────────────────────────────────────

def clear_body(doc):
    body = doc.element.body
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def set_columns(section, num, space_twips=360):
    sectPr = section._sectPr
    cols = sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sectPr.append(cols)
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), str(space_twips))


def para(doc, text, style=None, align=None, bold=False, italic=False, size=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        if size:
            r.font.size = Pt(size)
    return p


def heading(doc, text, level=1):
    style = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}[level]
    para(doc, text, style=style)


def body(doc, text):
    para(doc, text, style="Body Text")


def figure(doc, path, caption, width_in=3.35):
    if not os.path.exists(path):
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width_in))
    para(doc, caption, style="figure caption")


def _add_borders(t):
    tbl = t._tbl
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), "000000")
        borders.append(e)
    tblPr = tbl.tblPr
    tblPr.append(borders)


def _set_col_widths(t, total_in=3.3):
    t.autofit = False
    t.allow_autofit = False
    tbl = t._tbl
    tblPr = tbl.tblPr
    total_tw = int(total_in * 1440)
    n = len(t.columns)
    colw = total_tw // n

    # Force fixed layout so the renderer respects our widths instead of fitting
    # the table to its natural content width (which overflows the IEEE column).
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:type"), "dxa")
    tblW.set(qn("w:w"), str(total_tw))
    tblPr.append(tblW)

    grid = tbl.find(qn("w:tblGrid"))
    if grid is not None:
        for gc in grid.findall(qn("w:gridCol")):
            gc.set(qn("w:w"), str(colw))
    for row in t.rows:
        for cell in row.cells:
            cell.width = Pt(colw / 20.0)


def table(doc, header, rows, caption=None):
    if caption:
        para(doc, caption, style="table head")
    t = doc.add_table(rows=1, cols=len(header))
    _add_borders(t)
    for j, h in enumerate(header):
        cell = t.rows[0].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(7.5)
    for row in rows:
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = ""
            r = cells[j].paragraphs[0].add_run(str(val))
            r.font.size = Pt(7.5)
    _set_col_widths(t)
    return t


# ── build ───────────────────────────────────────────────────────────────────

def main():
    doc = Document(TEMPLATE)
    clear_body(doc)

    # Title block (section 1, full width)
    set_columns(doc.sections[0], 1)
    para(doc,
         "Adaptive Packet Routing with Deep Reinforcement Learning: "
         "Topology Generalization, Multi-Objective Trade-offs, Robustness, "
         "and Continual Adaptation",
         style="paper title")
    para(doc, "Vaibhav Rathod", style="Author", align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Department of Computer Science\nrathodvaibhav401@gmail.com",
         style="Author", align=WD_ALIGN_PARAGRAPH.CENTER)

    # Switch to two columns for the body
    sec = doc.add_section(WD_SECTION.CONTINUOUS)
    set_columns(sec, 2)

    # Abstract
    abs = ("Reinforcement learning (RL) is a promising route to adaptive packet "
           "forwarding, but most studies report a single algorithm on a single "
           "topology under a single objective. We present a unified empirical "
           "study of learned routing on a 10-node network simulated with an "
           "M/M/1 queuing-delay model, autocorrelated congestion and stochastic "
           "link failures. Five learners (vanilla DQN, Rainbow DQN, a "
           "graph-neural-network DQN, tabular Q-Routing, and a continual-learning "
           "variant) are compared against Dijkstra, ECMP and random routing with "
           "bootstrap confidence intervals over multiple seeds. Beyond the "
           "headline comparison we probe four practical questions that aggregate "
           "metrics hide: (i) does a structure-aware GNN policy transfer "
           "zero-shot to larger unseen graphs? (ii) where does the "
           "latency-reliability trade-off lie? (iii) how do policies degrade "
           "under worst-case rather than random failures? and (iv) does an agent "
           "forget an old traffic regime when adapting to a new one? We find the "
           "GNN policy is the sole Pareto-optimal method (lowest delay and "
           "highest delivery), keeps near-optimal route quality on graphs up to "
           "5x its training size, and resists targeted attacks that collapse a "
           "vanilla DQN from 0.91 to 0.60 delivery. We further show that naive "
           "sequential training forgets catastrophically (delivery on the old "
           "task drops from 0.77 to 0.09), that experience rehearsal fully "
           "prevents it, and that Elastic Weight Consolidation does not transfer "
           "to this destination-conditioned setting. All code, tests and results "
           "are released.")
    para(doc, "Abstract—" + abs, style="Abstract")
    para(doc, "Keywords—reinforcement learning; packet routing; graph neural "
              "networks; multi-objective optimization; continual learning; "
              "network simulation", style="Keywords")

    # I. Introduction
    heading(doc, "Introduction")
    body(doc,
         "Routing decides the path each packet takes through a network. Classical "
         "protocols such as OSPF compute shortest paths over a link-state metric "
         "(Dijkstra) and ECMP balances load over equal-cost paths. These methods "
         "are robust and well understood, but they react to congestion and "
         "failures only through periodic metric updates and do not learn from "
         "experience. Reinforcement learning offers an alternative: an agent "
         "observes link state and learns a forwarding policy that optimizes a "
         "long-horizon objective such as end-to-end delay and delivery.")
    body(doc,
         "The RL-routing literature, however, is fragmented. Papers typically "
         "demonstrate one algorithm, on one topology, under one scalarized "
         "objective, and report a single aggregate delivery or delay number. "
         "Several questions that matter for deployment are rarely answered "
         "together: whether a learned policy generalizes beyond its training "
         "graph; where it sits on the latency-reliability trade-off; how it "
         "behaves under adversarial rather than random failures; whether it can "
         "keep learning new regimes without forgetting old ones; and whether the "
         "aggregate metric is hiding per-flow starvation.")
    body(doc,
         "This paper makes these questions concrete on a common simulator and a "
         "common set of agents. Our contributions are: (1) a reproducible, "
         "multi-seed comparison of five learners and three classical baselines "
         "with statistical confidence intervals; (2) a zero-shot scalability "
         "study showing a graph-neural-network (GNN) policy transfers route "
         "quality to graphs up to five times larger than its training graph; "
         "(3) a multi-objective (Pareto) view, both across algorithms and within "
         "a single learner by sweeping its reward weights; (4) an "
         "adversarial-robustness study contrasting worst-case and average-case "
         "link failures; and (5) a continual-learning study quantifying "
         "catastrophic forgetting in learned routing and two remedies. We release "
         "the full code base, a 73-case test suite, and all result artifacts.")

    # II. Related Work
    heading(doc, "Related Work")
    body(doc,
         "Q-Routing [1] introduced distributed, online RL for routing with "
         "per-node Q-tables. Deep value-based RL [2] replaced tables with neural "
         "networks; Dueling architectures [3], prioritized replay [4] and their "
         "combination in Rainbow [5] improved sample efficiency and stability. "
         "Graph neural networks [6] and network-specific models such as RouteNet "
         "[7] encode topology directly, enabling generalization across graphs. "
         "Reproducibility concerns in deep RL [8] and robust evaluation protocols "
         "[9] motivate multi-seed training and confidence intervals. Continual "
         "learning addresses catastrophic forgetting with weight-space methods "
         "such as Elastic Weight Consolidation (EWC) [10] and data-space rehearsal. "
         "Fairness across flows is classically measured with Jain's index [11], "
         "and queuing delay with the M/M/1 model [12]. Our work differs by "
         "evaluating these dimensions jointly, on one simulator, with an honest "
         "account of where learned routing helps and where it does not.")

    # III. System Model
    heading(doc, "System Model and Methods")
    heading(doc, "Environment", 2)
    body(doc,
         "We model the network as a graph of 10 nodes and 17 bidirectional links, "
         "each with a bandwidth and a base propagation delay. The agent forwards a "
         "packet hop-by-hop from a random source to a random destination. Per-link "
         "queuing delay follows the M/M/1 sojourn-time model [12], "
         "W = base_delay + tx/(1-rho), so delay rises sharply as utilization rho "
         "approaches one. Background congestion evolves as an AR(1) process with "
         "bursty injections; loss rate rises with utilization in a RED-like manner; "
         "and links fail and recover under a two-state Markov process. The "
         "observation is a 70-dimensional vector of four features per link "
         "(congestion, normalized delay, loss, up/down) plus the normalized "
         "current and destination node identifiers.")
    heading(doc, "Reward", 2)
    body(doc,
         "The per-step reward combines a delivery bonus, a per-hop delay penalty "
         "weighted by w_delay, a packet-drop penalty weighted by w_drop, and "
         "penalties for loops and invalid actions. Latency and reliability are "
         "competing objectives, so exposing (w_delay, w_drop) lets us trace a "
         "Pareto front rather than commit to one scalarization (Section V-C).")
    heading(doc, "Agents", 2)
    body(doc,
         "We evaluate vanilla DQN, Rainbow DQN (Dueling + prioritized replay + "
         "n-step + Double-DQN), a GNN-DQN that performs GraphSAGE-style message "
         "passing over the live link features and is therefore topology-agnostic, "
         "and tabular Q-Routing. Classical baselines are Dijkstra (OSPF analog), "
         "ECMP and random forwarding. All neural agents are trained for multiple "
         "seeds with curriculum-increasing difficulty; evaluation uses greedy "
         "policies. We report bootstrap confidence intervals and effect sizes [8], "
         "[9].")

    # IV/V Results
    heading(doc, "Headline Comparison")
    body(doc,
         "Table I reports packet delivery ratio (PDR) and end-to-end delay across "
         "link-failure rates (200 episodes per setting). The GNN-DQN is the "
         "strongest learner: it maintains PDR = 1.00 at every failure rate while "
         "achieving the lowest delay (about 9-10 ms), because it routes from graph "
         "structure and reroutes around failures without degradation. Rainbow is "
         "competitive (PDR 0.90-0.94); vanilla DQN trails the classical baselines "
         "on PDR, which motivates the architectural additions. Q-Routing, an "
         "online table method, adapts slowly and degrades under failures. We note "
         "that in this simulator failed links are penalized but still traversable, "
         "so PDR measures successful navigation within a hop budget; this is why "
         "well-trained policies saturate near 1.0 and why we report additional, "
         "more discriminating metrics below.")
    table(doc,
          ["Method", "PDR @0%", "PDR @60%", "Delay @60%"],
          [["GNN-DQN", "1.00", "1.00", "10.2"],
           ["Rainbow", "0.93", "0.90", "15.4"],
           ["DQN", "0.86", "0.83", "12.8"],
           ["Q-Routing", "0.54", "0.40", "37.1"],
           ["Dijkstra", "0.95", "0.91", "15.1"],
           ["ECMP", "0.95", "0.92", "12.6"],
           ["Random", "1.00", "0.86", "36.5"]],
          caption="TABLE I.  DELIVERY (PDR) AND DELAY (ms) VS LINK-FAILURE RATE")
    figure(doc, os.path.join(RESULTS, "evaluation_comparison.png"),
           "Fig. 1.  Seven-algorithm comparison across link-failure rates "
           "(PDR, delay, jitter, throughput).")

    heading(doc, "Zero-Shot Scalability of the GNN Policy")
    body(doc,
         "Because the GNN operates on fixed-width node/edge features rather than a "
         "fixed node count, its trained weights apply to any graph. We evaluate "
         "the 10-node-trained policy zero-shot on Barabasi-Albert scale-free "
         "graphs of 20, 30 and 50 nodes. On well-connected graphs PDR is "
         "uninformative (any router eventually delivers), so we report delay "
         "stretch: the policy's mean delivered-path delay divided by Dijkstra's "
         "optimal. As Table II shows, the GNN holds stretch near 1.1-1.2 even at "
         "50 nodes (five times its training size), while a random walk degrades "
         "from 3.2 to 10.7. The learned policy thus transfers route quality, not "
         "merely reachability.")
    table(doc,
          ["Nodes", "Edges", "GNN stretch", "Random stretch"],
          [["10", "17", "1.14", "3.24"],
           ["20", "36", "1.09", "5.81"],
           ["30", "56", "1.09", "7.58"],
           ["50", "96", "1.22", "10.71"]],
          caption="TABLE II.  ZERO-SHOT TRANSFER TO LARGER UNSEEN GRAPHS "
                  "(GNN PDR = 1.00 throughout; stretch = delay / optimal)")
    figure(doc, os.path.join(RESULTS, "scalability.png"),
           "Fig. 2.  Delay stretch and PDR vs graph size for the GNN trained on "
           "10 nodes and applied zero-shot up to 50 nodes.")

    heading(doc, "Multi-Objective Trade-offs")
    body(doc,
         "Latency and reliability cannot both be maximized; a single reported "
         "operating point hides the trade-off. Plotting all algorithms in "
         "(delay, PDR) space at 20% failure, the GNN-DQN is the sole "
         "Pareto-optimal point (delay 9.24 ms, PDR 1.00), dominating every "
         "baseline. We also trace the trade-off within a single learner by "
         "training a DQN at a sweep of drop-penalty weights w_drop and measuring "
         "the achieved (mean delay, mean drops) operating point; the non-dominated "
         "weights w_drop = 15 and 30 form the learned policy's own latency-"
         "reliability frontier.")
    figure(doc, os.path.join(RESULTS, "pareto.png"),
           "Fig. 3.  Cross-algorithm Pareto front (left) and the learned policy's "
           "own front obtained by sweeping the reward's drop-penalty weight (right).")

    heading(doc, "Adversarial Robustness")
    body(doc,
         "Random link failures are average-case; a correlated outage or an "
         "attacker removes the most critical links. We compare random failures "
         "with a targeted attack that disables the highest edge-betweenness links "
         "at the same budget. At a 20% budget the GNN keeps PDR = 1.00 under both, "
         "whereas vanilla DQN collapses from 0.91 under random failures to 0.60 "
         "under the targeted attack, an exposure invisible to random-failure "
         "testing. (Dijkstra, which routes only the active subgraph, is shown as a "
         "hard-cut reachability reference rather than a competing policy.)")
    figure(doc, os.path.join(RESULTS, "adversarial.png"),
           "Fig. 4.  Delivery under random vs targeted (critical-link) failures.")

    heading(doc, "Per-Flow Fairness")
    body(doc,
         "Aggregate PDR can hide starvation of individual source-destination "
         "flows. Computing Jain's index [11] over per-flow delivery at 20% failure, "
         "the GNN, Dijkstra and ECMP are perfectly fair (index 1.0, minimum "
         "per-flow PDR 1.0), whereas vanilla DQN and Q-Routing starve some flows "
         "(minimum per-flow PDR 0.0) despite high means (Jain 0.92 each). Fairness "
         "is therefore a distinct axis from mean delivery.")

    heading(doc, "Continual Adaptation and Catastrophic Forgetting")
    body(doc,
         "A deployed agent must keep learning as demands change, but training on a "
         "new regime can overwrite the policy for the old one. We train a single "
         "DQN sequentially on two conflicting destination tasks on the same "
         "topology (deliver to nodes {7,8,9}, then to {0,1,2}) and measure the "
         "drop in delivery on the first task after learning the second (four "
         "seeds). Table III shows naive sequential training forgets "
         "catastrophically: Task-A delivery collapses from 0.77 to 0.09. "
         "Experience rehearsal, which simply retains old transitions in the replay "
         "buffer, fully prevents forgetting (it even improves slightly). EWC [10], "
         "the standard supervised continual-learning remedy, does not transfer "
         "here: because the policy is conditioned on the destination as an input "
         "feature, both tasks share the same weights, and a diagonal-Fisher anchor "
         "cannot isolate task-specific computation, so it performs no better than "
         "naive. The practical takeaway is that data-space rehearsal, nearly free "
         "in RL via the replay buffer, is the robust remedy.")
    table(doc,
          ["Method", "PDR A|A", "PDR A|B", "Forgetting"],
          [["Naive", "0.77", "0.09", "+0.67"],
           ["Rehearsal", "0.80", "0.83", "-0.03"],
           ["EWC", "0.82", "0.11", "+0.71"]],
          caption="TABLE III.  CATASTROPHIC FORGETTING AND REMEDIES. A|A = "
                  "Task-A PDR after training A; A|B = after training B.")
    figure(doc, os.path.join(RESULTS, "continual_learning.png"),
           "Fig. 5.  Per-task delivery after sequential training (left) and "
           "forgetting on the old task (right).")

    # VI. Discussion / limitations
    heading(doc, "Discussion and Limitations")
    body(doc,
         "Three honest caveats frame these results. First, the simulator models "
         "failed links as high-penalty but traversable, so PDR saturates and the "
         "more discriminating metrics (delay stretch, targeted-failure delivery, "
         "fairness, forgetting) carry the analysis. Second, most experiments use a "
         "single 10-node topology; the scalability study addresses size but only "
         "for the GNN and without failures. Third, all results are in simulation; "
         "a localhost data-plane prototype forwards real packets and reroutes on "
         "failure, but a hardware testbed remains future work. These do not affect "
         "the qualitative findings, but they bound the strength of the claims.")

    # VII. Conclusion
    heading(doc, "Conclusion")
    body(doc,
         "We presented a unified, reproducible study of RL-based routing that "
         "looks past a single aggregate number. A structure-aware GNN policy is "
         "Pareto-optimal, transfers zero-shot to graphs five times larger, and "
         "resists targeted attacks that break a vanilla DQN. We further quantified "
         "catastrophic forgetting in learned routing and showed that simple "
         "experience rehearsal is a robust remedy where EWC is not. The code, "
         "tests and results are released to support reproduction and extension to "
         "larger and hardware testbeds.")

    # References
    heading(doc, "References")
    refs = [
        "J. A. Boyan and M. L. Littman, “Packet routing in dynamically "
        "changing networks: A reinforcement learning approach,” in Adv. "
        "Neural Inf. Process. Syst. (NeurIPS), 1994.",
        "V. Mnih et al., “Human-level control through deep reinforcement "
        "learning,” Nature, vol. 518, pp. 529–533, 2015.",
        "Z. Wang et al., “Dueling network architectures for deep "
        "reinforcement learning,” in Proc. ICML, 2016.",
        "T. Schaul, J. Quan, I. Antonoglou, and D. Silver, “Prioritized "
        "experience replay,” in Proc. ICLR, 2016.",
        "M. Hessel et al., “Rainbow: Combining improvements in deep "
        "reinforcement learning,” in Proc. AAAI, 2018.",
        "W. L. Hamilton, R. Ying, and J. Leskovec, “Inductive "
        "representation learning on large graphs,” in Adv. Neural Inf. "
        "Process. Syst. (NeurIPS), 2017.",
        "K. Rusek et al., “RouteNet: Leveraging graph neural networks for "
        "network modeling and optimization in SDN,” IEEE J. Sel. Areas "
        "Commun., 2020.",
        "P. Henderson et al., “Deep reinforcement learning that matters,” "
        "in Proc. AAAI, 2018.",
        "R. Agarwal et al., “Deep reinforcement learning at the edge of the "
        "statistical precipice,” in Adv. Neural Inf. Process. Syst. "
        "(NeurIPS), 2021.",
        "J. Kirkpatrick et al., “Overcoming catastrophic forgetting in neural "
        "networks,” Proc. Natl. Acad. Sci. (PNAS), vol. 114, no. 13, 2017.",
        "R. Jain, D. Chiu, and W. Hawe, “A quantitative measure of fairness "
        "and discrimination for resource allocation in shared systems,” DEC "
        "Research Report TR-301, 1984.",
        "L. Kleinrock, Queueing Systems, Volume 1: Theory. Wiley, 1975.",
    ]
    for i, r in enumerate(refs, 1):
        para(doc, f"[{i}]\t{r}", style="references")

    doc.save(OUT)
    print(f"Paper written -> {OUT}")


if __name__ == "__main__":
    main()
